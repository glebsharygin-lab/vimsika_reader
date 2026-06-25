from __future__ import annotations

import argparse
import copy
import json
import re
import unicodedata
from pathlib import Path

from docx import Document
from pypdf import PdfReader


VERSE_NUMBERS = list(range(1, 23))
ROMAN = {
    1: "I",
    2: "II",
    3: "III",
    4: "IV",
    5: "V",
    6: "VI",
    7: "VII",
    8: "VIII",
    9: "IX",
    10: "X",
    11: "XI",
    12: "XII",
    13: "XIII",
    14: "XIV",
    15: "XV",
    16: "XVI",
    17: "XVII",
    18: "XVIII",
    19: "XIX",
    20: "XX",
    21: "XXI",
    22: "XXII",
}


def is_cjk(character: str) -> bool:
    codepoint = ord(character)
    return (
        0x3400 <= codepoint <= 0x4DBF
        or 0x4E00 <= codepoint <= 0x9FFF
        or 0xF900 <= codepoint <= 0xFAFF
    )


def is_word_character(character: str) -> bool:
    return unicodedata.category(character)[0] in {"L", "M", "N"}


def tokenize_text(
    text: str,
    passage_id: str,
    source_id: str,
) -> list[dict[str, object]]:
    tokens: list[dict[str, object]] = []
    index = 0
    while index < len(text):
        character = text[index]
        if is_cjk(character):
            end = index + 1
            token_type = "character"
        elif is_word_character(character):
            end = index + 1
            while end < len(text):
                next_character = text[end]
                if is_word_character(next_character):
                    end += 1
                    continue
                if (
                    next_character in {"'", "’", "-", "‐", "‑"}
                    and end + 1 < len(text)
                    and is_word_character(text[end + 1])
                ):
                    end += 1
                    continue
                break
            token_type = "syllable" if source_id == "tib_derge" else "word"
        else:
            index += 1
            continue

        token_number = len(tokens) + 1
        tokens.append(
            {
                "id": f"{passage_id}-{source_id}-t{token_number:05d}",
                "text": text[index:end],
                "start": index,
                "end": end,
                "type": token_type,
            }
        )
        index = end
    return tokens


def alignment_token_ids(
    text: str,
    tokens: list[dict[str, object]],
    target: str,
) -> list[str]:
    start = text.casefold().find(target.casefold())
    if start < 0:
        return []
    end = start + len(target)
    return [
        str(token["id"])
        for token in tokens
        if int(token["start"]) < end and int(token["end"]) > start
    ]


def is_alignable_token(
    token: dict[str, object],
    source_id: str,
) -> bool:
    text = str(token["text"])
    if not any(
        unicodedata.category(character)[0] in {"L", "M"}
        for character in text
    ):
        return False
    if text.casefold() == "vvs":
        return False
    if (
        text.isascii()
        and text.isupper()
        and (len(text) == 1 or re.fullmatch(r"[IVXLCDM]+", text))
    ):
        return False
    return True


def projected_token_span(
    target_tokens: list[dict[str, object]],
    source_index: int,
    source_count: int,
) -> list[str]:
    if not target_tokens or not source_count:
        return []
    target_count = len(target_tokens)
    start = round(source_index * target_count / source_count)
    end = round((source_index + 1) * target_count / source_count)
    start = min(start, target_count - 1)
    end = max(start + 1, min(end, target_count))
    return [str(token["id"]) for token in target_tokens[start:end]]


def candidate_alignments(
    passages: list[dict[str, object]],
    sources: list[dict[str, object]],
    sentence_alignments: list[dict[str, object]] | None = None,
) -> list[dict[str, object]]:
    candidates: list[dict[str, object]] = []
    target_source_ids = [
        str(source["id"])
        for source in sources
        if source["id"] != "san_levi_1925"
    ]
    alignments_by_verse: dict[int, list[dict[str, object]]] = {}
    for alignment in sentence_alignments or []:
        if (
            alignment.get("level") == "sentence"
            and alignment.get("status") == "dharmanexus-authorized"
        ):
            alignments_by_verse.setdefault(
                int(alignment["verse"]),
                [],
            ).append(alignment)
    for alignments in alignments_by_verse.values():
        alignments.sort(key=lambda item: int(item.get("order", 0)))

    for passage in passages:
        sanskrit_witness = passage["texts"]["san_levi_1925"]
        sanskrit_tokens = sanskrit_witness["tokens"]
        lexical_tokens = [
            token
            for token in sanskrit_tokens
            if is_alignable_token(token, "san_levi_1925")
        ]
        lexical_count = len(lexical_tokens)
        passage_alignments = alignments_by_verse.get(
            int(passage["number"]),
            [],
        )
        covered_token_ids: set[str] = set()

        def append_candidate(
            sanskrit_token: dict[str, object],
            source_index: int,
            source_count: int,
            target_tokens_by_source: dict[str, list[dict[str, object]]],
            source_alignment: dict[str, object] | None = None,
        ) -> None:
            target_token_ids: dict[str, list[str]] = {
                "san_levi_1925": [str(sanskrit_token["id"])]
            }
            target_methods: dict[str, str] = {
                "san_levi_1925": "source-token",
            }
            anchored_sources = set(
                source_alignment.get("anchoredSources", [])
                if source_alignment
                else []
            )
            for source_id in target_source_ids:
                target_tokens = target_tokens_by_source.get(source_id, [])
                if target_tokens:
                    target_token_ids[source_id] = projected_token_span(
                        target_tokens,
                        source_index,
                        source_count,
                    )
                    target_methods[source_id] = (
                        "dharmanexus-anchored-span-projection"
                        if source_id in anchored_sources
                        else "dharmanexus-segment-bounded-projection"
                    )
            candidates.append(
                {
                    "id": f"candidate-{sanskrit_token['id']}",
                    "verse": passage["number"],
                    "level": "token-span",
                    "status": "machine-suggested",
                    "confidence": "low",
                    "method": (
                        "dharmanexus-segment-bounded-token-projection-v1"
                        if source_alignment
                        else "monotonic-proportional-projection-v1"
                    ),
                    "label": str(sanskrit_token["text"]),
                    "note": (
                        "Automatically projected within a DharmaNexus-aligned segment. "
                        "This is a review candidate, not a philological assertion."
                        if source_alignment
                        else (
                            "Automatically projected by relative token position within "
                            "the passage. This is a review candidate, not a philological "
                            "assertion."
                        )
                    ),
                    "sourceAlignmentId": (
                        source_alignment.get("id")
                        if source_alignment
                        else None
                    ),
                    "targetMethods": target_methods,
                    "targetTokenIds": target_token_ids,
                }
            )

        for alignment in passage_alignments:
            alignment_target_ids = alignment.get("targetTokenIds", {})
            source_ids = set(
                alignment_target_ids.get("san_levi_1925", [])
            )
            source_tokens = [
                token
                for token in lexical_tokens
                if str(token["id"]) in source_ids
            ]
            if not source_tokens:
                continue
            target_tokens_by_source: dict[str, list[dict[str, object]]] = {}
            for source_id in target_source_ids:
                target_ids = set(alignment_target_ids.get(source_id, []))
                if not target_ids:
                    continue
                target_tokens_by_source[source_id] = [
                    token
                    for token in passage["texts"][source_id]["tokens"]
                    if str(token["id"]) in target_ids
                    and is_alignable_token(token, source_id)
                ]
            for source_index, sanskrit_token in enumerate(source_tokens):
                token_id = str(sanskrit_token["id"])
                if token_id in covered_token_ids:
                    continue
                covered_token_ids.add(token_id)
                append_candidate(
                    sanskrit_token,
                    source_index,
                    len(source_tokens),
                    target_tokens_by_source,
                    alignment,
                )

        fallback_targets = {
            source_id: [
                token
                for token in passage["texts"][source_id]["tokens"]
                if is_alignable_token(token, source_id)
            ]
            for source_id in target_source_ids
        }
        for source_index, sanskrit_token in enumerate(lexical_tokens):
            if str(sanskrit_token["id"]) in covered_token_ids:
                continue
            append_candidate(
                sanskrit_token,
                source_index,
                lexical_count,
                fallback_targets,
            )
    return candidates


SANSKRIT_SENTENCE_BOUNDARY = re.compile(
    r"(?:\|+|[।॥]+|/+|\n[ \t]*\n+)"
)
SANSKRIT_APPARATUS_MARKER = re.compile(
    r"\(\s*Vvs[_\s-]*\d+\s*\)",
    re.IGNORECASE,
)


def sanskrit_sentence_spans(text: str) -> list[tuple[int, int]]:
    spans: list[tuple[int, int]] = []
    cursor = 0
    for match in SANSKRIT_SENTENCE_BOUNDARY.finditer(text):
        raw_start = cursor
        raw_end = match.end()
        cursor = raw_end
        segment = text[raw_start:raw_end]
        content = SANSKRIT_APPARATUS_MARKER.sub("", segment)
        content = re.sub(r"[\s|/।॥\d()._-]+", "", content)
        if not content:
            if spans and re.search(r"\d", segment):
                spans[-1] = (spans[-1][0], raw_end)
            continue

        leading_marker = SANSKRIT_APPARATUS_MARKER.match(
            text,
            raw_start,
        )
        if leading_marker:
            raw_start = leading_marker.end()
        while raw_start < raw_end and text[raw_start].isspace():
            raw_start += 1
        while raw_end > raw_start and text[raw_end - 1].isspace():
            raw_end -= 1
        if raw_start < raw_end:
            spans.append((raw_start, raw_end))

    if cursor < len(text):
        raw_start = cursor
        raw_end = len(text)
        segment = text[raw_start:raw_end]
        content = SANSKRIT_APPARATUS_MARKER.sub("", segment)
        content = re.sub(r"[\s|/।॥\d()._-]+", "", content)
        if content:
            leading_marker = SANSKRIT_APPARATUS_MARKER.match(
                text,
                raw_start,
            )
            if leading_marker:
                raw_start = leading_marker.end()
            while raw_start < raw_end and text[raw_start].isspace():
                raw_start += 1
            while raw_end > raw_start and text[raw_end - 1].isspace():
                raw_end -= 1
            if raw_start < raw_end:
                spans.append((raw_start, raw_end))
    return spans


def tokens_overlapping_span(
    tokens: list[dict[str, object]],
    start: int,
    end: int,
) -> list[dict[str, object]]:
    return [
        token
        for token in tokens
        if int(token["start"]) < end and int(token["end"]) > start
    ]


def projected_sentence_span(
    source_tokens: list[dict[str, object]],
    target_tokens: list[dict[str, object]],
    sentence_tokens: list[dict[str, object]],
) -> list[dict[str, object]]:
    source_lexical = [
        token
        for token in source_tokens
        if is_alignable_token(token, "san_levi_1925")
    ]
    target_lexical = [
        token
        for token in target_tokens
        if is_alignable_token(token, "")
    ]
    sentence_ids = {str(token["id"]) for token in sentence_tokens}
    sentence_indexes = [
        index
        for index, token in enumerate(source_lexical)
        if str(token["id"]) in sentence_ids
    ]
    if not source_lexical or not target_lexical or not sentence_indexes:
        return []

    source_count = len(source_lexical)
    target_count = len(target_lexical)
    source_start = min(sentence_indexes)
    source_end = max(sentence_indexes) + 1
    target_start = round(source_start * target_count / source_count)
    target_end = round(source_end * target_count / source_count)
    target_start = min(target_start, target_count - 1)
    target_end = max(target_start + 1, min(target_end, target_count))
    return target_lexical[target_start:target_end]


def sentence_units_for_passage(
    passage: dict[str, object],
    sources: list[dict[str, object]],
) -> list[dict[str, object]]:
    sanskrit_witness = passage["texts"]["san_levi_1925"]
    sanskrit_text = str(sanskrit_witness["text"])
    sanskrit_tokens = sanskrit_witness["tokens"]
    units: list[dict[str, object]] = []

    for order, (start, end) in enumerate(
        sanskrit_sentence_spans(sanskrit_text),
        start=1,
    ):
        sentence_tokens = tokens_overlapping_span(
            sanskrit_tokens,
            start,
            end,
        )
        number = f"{passage['number']}.{order}"
        target_token_ids: dict[str, list[str]] = {
            "san_levi_1925": [
                str(token["id"])
                for token in sentence_tokens
            ]
        }
        targets = {
            "san_levi_1925": sanskrit_text[start:end],
        }

        for source in sources:
            source_id = str(source["id"])
            if source_id == "san_levi_1925":
                continue
            target_witness = passage["texts"][source_id]
            projected = projected_sentence_span(
                sanskrit_tokens,
                target_witness["tokens"],
                sentence_tokens,
            )
            if not projected:
                continue
            target_token_ids[source_id] = [
                str(token["id"])
                for token in projected
            ]
            target_text = str(target_witness["text"])
            targets[source_id] = target_text[
                int(projected[0]["start"]):int(projected[-1]["end"])
            ]

        units.append(
            {
                "id": f"sentence-v{passage['number']}-{order:03d}",
                "verse": passage["number"],
                "order": order,
                "number": number,
                "level": "sentence",
                "status": "machine-segmented",
                "confidence": "low",
                "method": (
                    "sanskrit-danda-and-paragraph-boundaries"
                    "+monotonic-proportional-projection-v1"
                ),
                "label": f"Sentence {number}",
                "note": (
                    "The Sanskrit boundary follows a danda, vertical stroke, slash, "
                    "or source paragraph break. Other witness spans are provisional "
                    "positional projections and require scholarly review."
                ),
                "targets": targets,
                "targetTokenIds": target_token_ids,
                "sourceCharacterSpan": {
                    "start": start,
                    "end": end,
                },
            }
        )
    return units


def load_authorized_alignments(
    path: Path | None,
    passages: list[dict[str, object]],
) -> list[dict[str, object]]:
    if path is None:
        return []
    payload = json.loads(path.read_text(encoding="utf-8"))
    records = payload.get("alignments", [])
    if not isinstance(records, list):
        raise ValueError("Authorized alignment import must contain an alignments list.")

    passages_by_number = {
        int(passage["number"]): passage
        for passage in passages
    }
    imported: list[dict[str, object]] = []
    for index, record in enumerate(records, start=1):
        if not isinstance(record, dict):
            raise ValueError(f"Alignment {index} must be an object.")
        verse = int(record["verse"])
        passage = passages_by_number.get(verse)
        if passage is None:
            raise ValueError(f"Alignment {index} references unknown verse {verse}.")
        alignment = {
            **record,
            "id": record.get("id", f"authorized-v{verse}-{index:04d}"),
            "verse": verse,
            "order": record.get("order", index),
            "level": record.get("level", "sentence"),
            "status": record.get("status", "externally-authorized"),
            "confidence": record.get("confidence", "reviewed"),
            "relation": record.get("relation", "parallel"),
            "label": record.get("label", f"Authorized alignment {index}"),
            "note": record.get(
                "note",
                "Imported from an explicitly authorized external export.",
            ),
        }
        target_token_ids = alignment.get("targetTokenIds", {})
        if not target_token_ids and alignment.get("targets"):
            target_token_ids = {}
            for source_id, target in alignment["targets"].items():
                witness = passage["texts"].get(source_id)
                if witness is None:
                    raise ValueError(
                        f"Alignment {alignment['id']} references unknown source {source_id}."
                    )
                target_token_ids[source_id] = alignment_token_ids(
                    witness["text"],
                    witness["tokens"],
                    str(target),
                )
        alignment["targetTokenIds"] = target_token_ids
        imported.append(alignment)
    return imported
SANSKRIT_STARTS = {
    1: "vijñaptimātram evedam",
    2: "na deśakālaniyamaḥ",
    3: "deśādiniyamaḥ siddhaḥ",
    4: "svapnopaghātavat kṛtyakriyā",
    5: "tiraścāṃ saṃbhavaḥ",
    6: "yadi tatkarmabhis tatra",
    7: "karmaṇo vāsanānyatra",
    8: "rūpādyāyatanāstitvaṃ",
    9: "yataḥ svabījād vijñaptir",
    10: "tathā pudgalanairātmyapraveśo hi",
    11: "na tad ekaṃ na cānekaṃ",
    12: "ṣaṭkena yugapad yogāt",
    13: "paramāṇor asaṃyogāt",
    14: "digbhāgabhedī yasyāsti",
    15: "ekatve na krameṇetir",
    16: "pratyakṣabuddhiḥ svapnādau",
    17: "uktaṃ yathā tadābhāsā",
    18: "anyonyādhipatitvena",
    19: "maraṇam paravijñaptiviśeṣād",
    20: "kathaṃ vā daṇḍakāraṇyaśūnyatvam",
    21: "paracittavidāṃ jñānam",
    22: "vijñaptimātratāsiddhiḥ",
}
SANSKRIT_ROOTS = {
    1: (
        "vijñaptimātram evedam asadarthāvabhāsanāt |\n"
        "yadvat taimirikasyāsat keśoṇḍukādidarśanaṃ || 1 ||"
    ),
    2: (
        "na deśakālaniyamaḥ saṃtānāniyamo na ca |\n"
        "na ca kṛtyakriyā yuktā vijñaptir yadi nārthataḥ || 2 ||"
    ),
    3: (
        "deśādiniyamaḥ siddhaḥ svapnavat\n"
        "pretavat punaḥ |\n"
        "saṃtānāniyamaḥ\n"
        "sarvaiḥ pūyanadyādidarśane || 3 ||"
    ),
    4: (
        "svapnopaghātavat kṛtyakriyā\n"
        "narakavat punaḥ |\n"
        "sarvaṃ\n"
        "narakapālādidarśane taiś ca bādhane || 4 ||"
    ),
    5: (
        "tiraścāṃ saṃbhavaḥ svarge yathā na narake tathā |\n"
        "na pretānāṃ yatas tajjaṃ duḥkhaṃ nānubhavanti te || 5 ||"
    ),
    6: (
        "yadi tatkarmabhis tatra bhūtānāṃ saṃbhavas tathā |\n"
        "iṣyate pariṇāmaś ca kiṃ vijñānasya neṣyate || 6 ||"
    ),
    7: (
        "karmaṇo vāsanānyatra phalam anyatra kalpyate |\n"
        "tatraiva neṣyate yatra vāsanā kiṃ nu kāraṇaṃ || 7 ||"
    ),
    8: (
        "rūpādyāyatanāstitvaṃ tadvineyajanaṃ prati |\n"
        "abhiprāyavaśād uktam upapādukasattvavat || 8 ||"
    ),
    9: (
        "yataḥ svabījād vijñaptir yadābhāsā pravartate |\n"
        "dvividhāyatanatvena te tasyā munir abravīt || 9 ||"
    ),
    10: (
        "tathā pudgalanairātmyapraveśo hi\n"
        "anyathā punaḥ |\n"
        "deśanā dharmanairātmyapraveśaḥ\n"
        "kalpitātmanā || 10 ||"
    ),
    11: (
        "na tad ekaṃ na cānekaṃ viṣayaḥ paramāṇuśaḥ |\n"
        "na ca te saṃhatā yasmāt paramāṇur na sidhyati || 11 ||"
    ),
    12: (
        "ṣaṭkena yugapad yogāt paramāṇoḥ ṣaḍaṃśatā |\n"
        "ṣaṇṇāṃ samānadeśatvāt piṇḍaḥ syād aṇumātrakaḥ || 12 ||"
    ),
    13: (
        "paramāṇor asaṃyogāt tatsaṃghāte 'sti kasya saḥ |\n"
        "na cānavayavatvena tatsaṃyogo na sidhyati || 13 ||"
    ),
    14: (
        "digbhāgabhedī yasyāsti tasyaikatvaṃ na yujyate |\n"
        "chāyāvṛtī kathaṃ vā\n"
        "anyo na piṇḍaś cen na tasya te || 14 ||"
    ),
    15: (
        "ekatve na krameṇetir yugapan na grahāgrahau |\n"
        "vicchinnānekavṛttiś ca sūkṣmānīkṣā ca no bhavet || 15 ||"
    ),
    16: (
        "pratyakṣabuddhiḥ svapnādau yathā\n"
        "sā ca yadā tadā |\n"
        "na so 'rtho dṛśyate tasya pratyakṣatvaṃ kathaṃ mataṃ || 16 ||"
    ),
    17: (
        "uktaṃ yathā tadābhāsā vijñaptiḥ\n"
        "smaraṇaṃ tataḥ |\n"
        "svapnadṛgviṣayābhāvaṃ nāprabuddho 'vigacchati || 17 ||"
    ),
    18: (
        "anyonyādhipatitvena vijñaptiniyamo mithaḥ |\n"
        "middhenopahataṃ cittaṃ svapne tenāsamaṃ phalaṃ || 18 ||"
    ),
    19: (
        "maraṇam paravijñaptiviśeṣād vikriyā yathā |\n"
        "smṛtilopādikānyeṣāṃ piśācādimanovaśāt || 19 ||"
    ),
    20: (
        "kathaṃ vā daṇḍakāraṇyaśūnyatvam ṛṣikopataḥ |\n"
        "manodaṇḍo mahāvadyaḥ kathaṃ vā tena sidhyati || 20 ||"
    ),
    21: (
        "paracittavidāṃ jñānam ayathārthaṃ kathaṃ yathā |\n"
        "svacittajñānaṃ\n"
        "ajñānād yathā buddhasya gocaraḥ || 21 ||"
    ),
    22: (
        "vijñaptimātratāsiddhiḥ svaśaktisadṛśī mayā |\n"
        "kṛteyaṃ sarvathā sā tu na cintyā\n"
        "buddhagocaraḥ || 22 ||"
    ),
}


def read_docx_paragraphs(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs if paragraph.text.strip()]


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
    text = re.sub(r"(?<=\w)-\s*\n\s*(?=\w)", "", text)
    return re.sub(r"[ \t]+", " ", text)


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def slice_by_markers(
    text: str,
    markers: dict[int, re.Pattern[str]],
    first_start: int | None = None,
) -> dict[int, str]:
    positions: dict[int, int] = {}
    cursor = 0
    for number in VERSE_NUMBERS:
        match = markers[number].search(text, cursor)
        if not match:
            raise ValueError(f"Could not locate verse {number}")
        positions[number] = match.start()
        cursor = match.start() + 1

    if first_start is not None:
        positions[VERSE_NUMBERS[0]] = first_start

    result: dict[int, str] = {}
    for number in VERSE_NUMBERS:
        start = positions[number]
        end = positions.get(number + 1, len(text))
        result[number] = clean_text(text[start:end])
    return result


def extract_sanskrit(path: Path) -> dict[int, str]:
    text = "\n\n".join(read_docx_paragraphs(path))
    markers = {
        number: re.compile(re.escape(start), re.IGNORECASE)
        for number, start in SANSKRIT_STARTS.items()
    }
    return slice_by_markers(text, markers, first_start=0)


def extract_tibetan(path: Path) -> dict[int, str]:
    text = path.read_text(encoding="utf-8")
    markers = {
        number: re.compile(rf"(?m)^{ROMAN[number]}\s*$")
        for number in VERSE_NUMBERS
    }
    return slice_by_markers(text, markers, first_start=0)


def extract_english(path: Path) -> dict[int, str | None]:
    text = read_pdf(path)
    available_numbers = [1, 2, 3, 4, *range(11, 23)]
    markers = {
        number: re.compile(rf"\bVerse\s+{number}\.", re.IGNORECASE)
        for number in available_numbers
    }
    positions: dict[int, int] = {}
    cursor = 0
    for number in available_numbers:
        match = markers[number].search(text, cursor)
        if not match:
            raise ValueError(f"Could not locate English verse {number}")
        positions[number] = match.start()
        cursor = match.start() + 1

    statement_start = text.find("I. A Statement of the View")
    if statement_start >= 0:
        positions[1] = statement_start

    passages: dict[int, str | None] = {number: None for number in VERSE_NUMBERS}
    for index, number in enumerate(available_numbers):
        start = positions[number]
        end = (
            positions[available_numbers[index + 1]]
            if index + 1 < len(available_numbers)
            else len(text)
        )
        passages[number] = clean_text(text[start:end])
    return passages


def extract_german(path: Path) -> dict[int, str]:
    text = read_pdf(path)
    markers = {
        number: re.compile(rf"\bV\.\s*{number}\b", re.IGNORECASE)
        for number in VERSE_NUMBERS
    }
    opening = text.find("KAPITEL")
    return slice_by_markers(text, markers, first_start=max(opening, 0))


def extract_french(path: Path) -> dict[int, str]:
    text = clean_text("\n\n".join(read_docx_paragraphs(path)))
    starts = {
        1: "Le Grand Véhicule établit",
        2: "2. – Si la Notification",
        3: "3 a. La détermination",
        4: "4 a. Pour l'accomplissement",
        5: "5. S'il peut",
        6: "6. Si vous admettez",
        7: "7. Vous imaginez",
        8: "8. S'il est parlé",
        9: "Le germe propre d'où la Notification",
        10: "C'est que l'entrée dans le Sans-Soi",
        11: "11. Le Domaine",
        12: "12 ab. Si l'Atome",
        13: "13 ab. – Si l'Atome",
        14: "14 ab. Du moment",
        15: "15. Dans",
        16: "16 a. L'idée",
        17: "17 a. Nous avons montré",
        18: "18 ab. C'est par",
        19: "19. La mort",
        20: "20 ab. – Ou bien",
        21: "21 abc. Chez ceux",
        22: "22 abc. Cette démonstration",
    }
    positions: dict[int, int] = {}
    cursor = 0
    for number in VERSE_NUMBERS:
        position = text.find(starts[number], cursor)
        if position < 0:
            raise ValueError(f"Could not locate French verse {number}: {starts[number]}")
        positions[number] = position
        cursor = position + 1

    return {
        number: clean_text(
            text[positions[number] : positions.get(number + 1, len(text))]
        )
        for number in VERSE_NUMBERS
    }


def cell_has_verse(cell: str, number: int) -> bool:
    return bool(re.search(rf"\(\s*{number}(?:[a-d-]*)?\s*\)", cell, re.IGNORECASE))


def extract_chinese_table(path: Path) -> dict[str, dict[int, str]]:
    document = Document(path)
    table = document.tables[0]
    rows = [[clean_text(cell.text) for cell in row.cells] for row in table.rows]
    columns = {
        "zho_xuanzang": 0,
        "zho_paramartha": 1,
        "zho_bodhiruci": 2,
    }

    starts: dict[int, int] = {}
    cursor = 0
    for number in VERSE_NUMBERS:
        for row_index in range(cursor, len(rows)):
            if cell_has_verse(rows[row_index][3], number):
                starts[number] = row_index
                cursor = row_index + 1
                break
        else:
            raise ValueError(f"Could not locate Chinese table verse {number}")
    starts[1] = 1

    passages = {version_id: {} for version_id in columns}
    for number in VERSE_NUMBERS:
        start = starts[number]
        end = starts.get(number + 1, len(rows))
        passage_rows = rows[start:end]
        for version_id, column_index in columns.items():
            values = [row[column_index] for row in passage_rows if row[column_index]]
            passages[version_id][number] = clean_text("\n\n".join(values))

    return passages


def load_segmented_witness_records(
    path: Path,
    require_text: bool = True,
) -> dict[int, dict[str, str]]:
    records = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(records, list):
        raise ValueError(f"Segmented witness must be a list: {path}")
    passages = {}
    for record in records:
        verse = int(record["verse"])
        passages[verse] = {
            "text": clean_text(str(record.get("text", ""))),
            "status": str(record.get("status", "machine-segmented")),
            "note": str(record.get("note", "")),
            "method": str(record.get("method", "")),
        }
    missing = [
        number
        for number in VERSE_NUMBERS
        if number not in passages
        or (require_text and not passages[number]["text"])
    ]
    if missing:
        raise ValueError(
            f"Segmented witness {path} is missing verses: {missing}"
        )
    return passages


def load_segmented_witness(path: Path) -> dict[int, str]:
    records = load_segmented_witness_records(path)
    return {
        number: record["text"]
        for number, record in records.items()
    }


def source_records() -> list[dict[str, object]]:
    sources = [
        {
            "id": "san_levi_1925",
            "label": "Sanskrit · Lévi 1925",
            "shortLabel": "Sanskrit",
            "language": "Sanskrit",
            "languageCode": "san",
            "script": "Latin transliteration",
            "role": "edition",
            "color": "#a64b2a",
            "citation": "Sylvain Lévi (ed.), Vijñaptimātratāsiddhi, Paris, 1925.",
            "file": "san_levi.docx",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "DOCX text layer; passage boundaries inferred from verse openings.",
        },
        {
            "id": "san_silk_2016",
            "label": "Sanskrit · Silk 2016",
            "shortLabel": "Silk Sanskrit",
            "language": "Sanskrit",
            "languageCode": "san",
            "script": "Latin transliteration",
            "role": "edition",
            "color": "#b85f3f",
            "citation": (
                "Jonathan A. Silk, Materials Toward the Study of "
                "Vasubandhu’s Viṃśikā (I), Harvard Oriental Series 81, "
                "2016; 2018 open-access edition."
            ),
            "file": "view.pdf, PDF pages 215–225",
            "rights": "cc-by-sa",
            "rightsLabel": "CC BY-SA",
            "extraction": (
                "Sanskrit reading text segmented from Roman-numbered sections. "
                "Verse 1 is partial in the supplied published reading text."
            ),
        },
        {
            "id": "san_tola_dragonetti_2004",
            "label": "Sanskrit · Tola & Dragonetti 2004",
            "shortLabel": "Tola–Dragonetti Sanskrit",
            "language": "Sanskrit",
            "languageCode": "san",
            "script": "Latin transliteration",
            "role": "edition",
            "color": "#8e6541",
            "citation": (
                "Fernando Tola and Carmen Dragonetti, Being as Consciousness: "
                "Yogācāra Philosophy of Buddhism, Motilal Banarsidass, 2004."
            ),
            "file": "Being as Consciousness, PDF pages 165–175",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "Legacy Latin text layer segmented from Sanskrit kārikā openings; "
                "diacritics and printed note markers require proofing."
            ),
        },
        {
            "id": "san_ruzsa_szegedi_2015",
            "label": "Sanskrit · Ruzsa & Szegedi 2015",
            "shortLabel": "Ruzsa–Szegedi",
            "language": "Sanskrit",
            "languageCode": "san",
            "script": "Latin transliteration",
            "role": "critical-edition",
            "color": "#9d4938",
            "citation": (
                "Ferenc Ruzsa and Mónika Szegedi, “Vasubandhu’s Viṁśikā: "
                "A Critical Edition,” Távol-keleti Tanulmányok 2015/1, 127–158."
            ),
            "file": "Vasubandhu_Visikajanak_kritikai_kiadasa.pdf, PDF pages 8–31",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "Critical reading text segmented from numbered kārikā headings. "
                "The full extracted apparatus remains in source-witnesses."
            ),
        },
        {
            "id": "san_balcerowicz_nowakowska_1999",
            "label": "Sanskrit · Balcerowicz & Nowakowska 1999",
            "shortLabel": "Balcerowicz Sanskrit",
            "language": "Sanskrit",
            "languageCode": "san",
            "script": "Devanagari",
            "role": "edition-draft",
            "color": "#72564b",
            "citation": (
                "Piotr Balcerowicz and Monika Nowakowska, Studia "
                "Indologiczne 6 (1999), 5–44."
            ),
            "file": "si06(1999).pdf, PDF pages 6–17",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "Rule-based Unicode Devanagari draft converted from the "
                "embedded AmritaA font encoding. The draft requires scholarly "
                "proofing before final collation."
            ),
        },
        {
            "id": "san_tiwari_1995",
            "label": "Sanskrit · Tiwari 1995",
            "shortLabel": "Tiwari Sanskrit",
            "language": "Sanskrit",
            "languageCode": "san",
            "script": "Devanagari",
            "role": "edition",
            "color": "#a34f35",
            "citation": (
                "Mahesh Tiwari, ed. and trans., Vijñaptimātratāsiddhiḥ: "
                "Viṃśatikā-Triṃśikābhidhāna-prakaraṇadvayātmikā, "
                "Chaukhambha Vidyabhavan, Varanasi, 1995."
            ),
            "file": (
                "Internet Archive item vijnaptimatrasiddhiofvasubandhu"
                "vigyanptimaheshtiwarichowkambhasanskritseries_492_y, "
                "printed pages 25–26"
            ),
            "rights": "rights-review",
            "rightsLabel": "Rights status requires verification",
            "extraction": (
                "Numbered Sanskrit kārikā appendix extracted from DjVu OCR. "
                "The Devanagari text is preliminary and requires proofing "
                "against the scan."
            ),
        },
        {
            "id": "san_kalupahana_1987",
            "label": "Sanskrit · Kalupahana 1987",
            "shortLabel": "Kalupahana Sanskrit",
            "language": "Sanskrit",
            "languageCode": "san",
            "script": "Latin transliteration",
            "role": "edition",
            "color": "#9f5b45",
            "citation": (
                "David J. Kalupahana, The Principles of Buddhist Psychology, "
                "SUNY Press, 1987, Appendix II."
            ),
            "file": "The Principles of Buddhist Psychology, PDF pages 184–203",
            "rights": "rights-review",
            "rightsLabel": "Rights status requires verification / OCR proofing needed",
            "extraction": (
                "Appendix II Sanskrit text extracted from the PDF text layer. "
                "The romanization is noisy and requires proofing against the "
                "printed appendix."
            ),
        },
        {
            "id": "tib_derge",
            "label": "Tibetan · Derge",
            "shortLabel": "Tibetan",
            "language": "Tibetan",
            "languageCode": "bod",
            "script": "Wylie transliteration",
            "role": "canonical-translation",
            "color": "#8b6a20",
            "citation": "Nyi shu pa’i ’grel pa, Derge Bstan ’gyur tradition.",
            "file": "tib_derge.txt",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "UTF-8 text; sections inferred from Roman-numeral headings.",
        },
        {
            "id": "tib_silk_dunhuang_2017",
            "label": "Tibetan · Dunhuang PT 797 (Silk 2017)",
            "shortLabel": "Dunhuang PT 797",
            "language": "Tibetan",
            "languageCode": "bod",
            "script": "Latin transliteration",
            "role": "manuscript-transcription",
            "color": "#7a7048",
            "citation": (
                "Jonathan A. Silk, “Materials Toward the Study of Vasubandhu’s "
                "Viṃśikā (II): An edition of the Dunhuang Manuscript Pelliot "
                "tibétain 797,” Revue d’Etudes Tibétaines 39 (2017), 342–360."
            ),
            "file": "Materials_Toward_the_Study_of_Vasubandhu (1).pdf, pages 344–358",
            "rights": "open-access",
            "rightsLabel": "Open access / research use",
            "extraction": (
                "Pelliot tibétain 797 critical transcription extracted from "
                "Silk's imposed Roman divisions. Verse rubrics and line notes "
                "require scholarly review."
            ),
        },
        {
            "id": "tib_lvp_1911",
            "label": "Tibetan · La Vallée-Poussin 1911",
            "shortLabel": "LVP Tibetan",
            "language": "Tibetan",
            "languageCode": "bod",
            "script": "Latin transliteration",
            "role": "edition-ocr",
            "color": "#716238",
            "citation": (
                "Louis de La Vallée-Poussin, “Vasubandhu: "
                "Viṃśakakārikāprakaraṇa. Traité des vingt ślokas avec le "
                "commentaire de l’auteur,” Le Muséon 30 (1911), 53–90."
            ),
            "file": "lemuson30soc1i (1).pdf, pages 54–65",
            "rights": "public-domain",
            "rightsLabel": "Public domain / OCR proofing needed",
            "extraction": (
                "Le Muséon Tibetan OCR layer segmented by visible verse "
                "anchors. OCR and segmentation require line-by-line correction."
            ),
        },
        {
            "id": "zho_xuanzang",
            "label": "Chinese · Xuanzang",
            "shortLabel": "Xuanzang",
            "language": "Chinese",
            "languageCode": "zho",
            "script": "Traditional Chinese",
            "role": "canonical-translation",
            "color": "#9a3f57",
            "citation": "唯識二十論, translated by Xuanzang, Taishō T1590.",
            "file": "zho_xuangzan.docx",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "First column of the supplied four-column comparative table.",
        },
        {
            "id": "zho_hamilton_xuanzang_1938",
            "label": "Chinese · Xuanzang / Hamilton 1938",
            "shortLabel": "Hamilton Chinese",
            "language": "Chinese",
            "languageCode": "zho",
            "script": "Traditional Chinese",
            "role": "edition-ocr",
            "color": "#ad5870",
            "citation": (
                "Facing Chinese text printed with Clarence H. Hamilton, trans., "
                "“Wei-shih-er-shih-lun or The Treatise in Twenty Stanzas on "
                "Representation-Only by Vasubandhu,” 1938; Hamilton identifies "
                "the Chinese base as Xuanzang’s text established by the Chinese "
                "Academy of Buddhist Learning, Nanking, 1930."
            ),
            "file": "Hamilton 1938 scan, facing Chinese pages",
            "rights": "public-domain",
            "rightsLabel": "Public domain / OCR proofing needed",
            "extraction": (
                "Vertical Chinese scan reconstructed with column-aware Windows "
                "OCR and segmented by base-assisted collation against the clean "
                "Xuanzang witness. OCR and passage boundaries require proofing."
            ),
        },
        {
            "id": "zho_paramartha",
            "label": "Chinese · Paramārtha",
            "shortLabel": "Paramārtha",
            "language": "Chinese",
            "languageCode": "zho",
            "script": "Traditional Chinese",
            "role": "canonical-translation",
            "color": "#7f4a77",
            "citation": "大乘唯識論, translated by Paramārtha, Taishō T1589.",
            "file": "zho_xuangzan.docx",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "Second column of the supplied four-column comparative table.",
        },
        {
            "id": "zho_bodhiruci",
            "label": "Chinese · Prajñāruci",
            "shortLabel": "Prajñāruci",
            "language": "Chinese",
            "languageCode": "zho",
            "script": "Traditional Chinese",
            "role": "canonical-translation",
            "color": "#6d5a8d",
            "citation": "唯識論, translated by Prajñāruci, Taishō T1588.",
            "file": "zho_xuangzan.docx",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "Third column of the supplied four-column comparative table.",
        },
        {
            "id": "eng_das",
            "label": "English · Nilanjan Das",
            "shortLabel": "Das",
            "language": "English",
            "languageCode": "eng",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#2d6f8f",
            "citation": "Nilanjan Das, Twenty Verses with Auto-Commentary, online draft.",
            "file": "eng_das.pdf",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "Searchable PDF text layer; passages inferred from verse labels.",
        },
        {
            "id": "eng_silk_2016",
            "label": "English · Silk 2016",
            "shortLabel": "Silk",
            "language": "English",
            "languageCode": "eng",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#355f91",
            "citation": (
                "Jonathan A. Silk, Materials Toward the Study of "
                "Vasubandhu’s Viṃśikā (I), Harvard Oriental Series 81, "
                "2016; 2018 open-access edition."
            ),
            "file": "view.pdf, PDF pages 229–250",
            "rights": "cc-by-sa",
            "rightsLabel": "CC BY-SA",
            "extraction": (
                "Searchable PDF text layer; passages segmented from the "
                "translator’s Roman-numbered kārikā sections."
            ),
        },
        {
            "id": "eng_siderits_2007",
            "label": "English · Siderits 2007",
            "shortLabel": "Siderits",
            "language": "English",
            "languageCode": "eng",
            "script": "Latin",
            "role": "modern-translation-excerpt",
            "color": "#487d7a",
            "citation": (
                "Mark Siderits, Buddhism as Philosophy: An Introduction, "
                "Ashgate, 2007, chapter 8."
            ),
            "file": "Buddhism as Philosophy, PDF chapter 8",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus / partial witness",
            "extraction": (
                "Selected printed Viṃśatikā passages extracted from the PDF "
                "text layer. Verses not quoted in the supplied chapter are "
                "left blank."
            ),
        },
        {
            "id": "eng_kalupahana_1987",
            "label": "English · Kalupahana 1987",
            "shortLabel": "Kalupahana",
            "language": "English",
            "languageCode": "eng",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#4f7396",
            "citation": (
                "David J. Kalupahana, The Principles of Buddhist Psychology, "
                "SUNY Press, 1987, Appendix II."
            ),
            "file": "The Principles of Buddhist Psychology, PDF pages 184–203",
            "rights": "rights-review",
            "rightsLabel": "Rights status requires verification / OCR proofing needed",
            "extraction": (
                "English translation and annotation extracted from Appendix II "
                "via the PDF text layer. OCR-like spelling and line-break noise "
                "require proofreading."
            ),
        },
        {
            "id": "eng_anacker_2005",
            "label": "English · Anacker 2005",
            "shortLabel": "Anacker",
            "language": "English",
            "languageCode": "eng",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#76518d",
            "citation": (
                "Stefan Anacker, Seven Works of Vasubandhu: The Buddhist "
                "Psychological Doctor, revised edition, Motilal "
                "Banarsidass, 2005."
            ),
            "file": "Seven Works of Vasubandhu, PDF pages 174–188",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "Searchable PDF text layer; passages segmented from the "
                "numbered kārikā translations. Printed notes are excluded."
            ),
        },
        {
            "id": "eng_tola_dragonetti_2004",
            "label": "English · Tola & Dragonetti 2004",
            "shortLabel": "Tola–Dragonetti",
            "language": "English",
            "languageCode": "eng",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#3f7182",
            "citation": (
                "Fernando Tola and Carmen Dragonetti, Being as "
                "Consciousness: Yogācāra Philosophy of Buddhism, "
                "Motilal Banarsidass, 2004."
            ),
            "file": "Being as Consciousness, PDF pages 176–195",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "Searchable PDF text layer; passages segmented from the "
                "numbered stanza translations."
            ),
        },
        {
            "id": "eng_kochumuttom_1982",
            "label": "English · Kochumuttom 1982",
            "shortLabel": "Kochumuttom",
            "language": "English",
            "languageCode": "eng",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#536e9c",
            "citation": (
                "Thomas A. Kochumuttom, A Buddhist Doctrine of Experience: "
                "A New Translation and Interpretation of the Works of Vasubandhu "
                "the Yogācārin, Motilal Banarsidass, 1982."
            ),
            "file": "A Buddhist Doctrine of Experience, PDF pages 25–32",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "Full English translation and auto-commentary segmented from "
                "numbered kārikā headings. OCR typography and Sanskrit forms "
                "require proofing."
            ),
        },
        {
            "id": "eng_wood_1991",
            "label": "English · Wood 1991",
            "shortLabel": "Wood",
            "language": "English",
            "languageCode": "eng",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#597c65",
            "citation": (
                "Thomas E. Wood, Mind Only: A Philosophical and Doctrinal "
                "Analysis of the Vijñānavāda, University of Hawai‘i Press, 1991."
            ),
            "file": "Mind Only, PDF pages 113–118",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "Verse-only English translation segmented from numbered stanzas. "
                "The printed Sanskrit and surrounding chapter analysis are excluded."
            ),
        },
        {
            "id": "eng_cook_1999",
            "label": "English · Cook 1999",
            "shortLabel": "Cook",
            "language": "English",
            "languageCode": "eng",
            "script": "Latin",
            "role": "modern-translation-from-chinese",
            "color": "#8a5e35",
            "citation": (
                "Francis H. Cook, trans., “The Treatise in Twenty Verses on "
                "Consciousness Only,” in Three Texts on Consciousness Only, "
                "Numata Center for Buddhist Translation and Research, 1999."
            ),
            "file": "vasubandhu_-_vimshatika__with_vs_comment___cook_.pdf, pages 7–24",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "English translation from Xuanzang’s Chinese T1590. Xuanzang’s "
                "twenty-one numbered verses are mapped by content to the shell’s "
                "twenty-two Sanskrit passages."
            ),
        },
        {
            "id": "eng_hamilton_1938",
            "label": "English · Hamilton 1938",
            "shortLabel": "Hamilton",
            "language": "English",
            "languageCode": "eng",
            "script": "Latin",
            "role": "modern-translation-from-chinese-ocr",
            "color": "#6e6f9e",
            "citation": (
                "Clarence H. Hamilton, trans., “Wei-shih-er-shih-lun or "
                "The Treatise in Twenty Stanzas on Representation-Only by "
                "Vasubandhu,” 1938."
            ),
            "file": "Hamilton 1938 scan, printed pages 19–79",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus / OCR proofing needed",
            "extraction": (
                "Scanned English translation from Xuanzang’s Chinese recovered "
                "with local Windows OCR and segmented by semantic stanza anchors. "
                "OCR typography, notes, and boundaries require proofing."
            ),
        },
        {
            "id": "eng_cronk_1998",
            "label": "English · Cronk 1998",
            "shortLabel": "Cronk",
            "language": "English",
            "languageCode": "eng",
            "script": "Latin",
            "role": "modern-adaptation",
            "color": "#6f7142",
            "citation": (
                "George Cronk, rendition and editing, “Twenty Verses on "
                "Consciousness-Only (Vimsatika-Karika),” 1998."
            ),
            "file": "twenty-verses-on-consciousness-only-vimsatika-karika.pdf",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "Adapted English rendition segmented from bracketed verse "
                "labels. The supplied PDF marks 17 of the 22 local verse "
                "passages."
            ),
        },
        {
            "id": "pol_balcerowicz_nowakowska_1999",
            "label": "Polish · Balcerowicz & Nowakowska 1999",
            "shortLabel": "Balcerowicz–Nowakowska",
            "language": "Polish",
            "languageCode": "pol",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#8a4968",
            "citation": (
                "Piotr Balcerowicz and Monika Nowakowska, “Wasubandhu: "
                "‘Dowód na wyłączne istnienie treści świadomości w "
                "dwudziestu strofach’,” Studia Indologiczne 6 (1999), 18–35."
            ),
            "file": "si06(1999).pdf, PDF pages 18–35",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "PolishTimes encoding recovered and segmented from numbered "
                "kārikā translations; printed footnotes are excluded."
            ),
        },
        {
            "id": "hun_szanyi_2015",
            "label": "Hungarian · Szanyi 2015",
            "shortLabel": "Szanyi",
            "language": "Hungarian",
            "languageCode": "hun",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#3f786d",
            "citation": (
                "Szilvia Szanyi, “Buddhista idealizmus: Vasubandhu Viṃśatikā "
                "című művének filozófiai elemzése,” Távol-keleti "
                "Tanulmányok 2015/2, 107–136."
            ),
            "file": "Buddhista_idealizmus_Vasubandhu_Visatik.pdf, PDF pages 4–26",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "CID font recovered to Unicode; translated verse quotations "
                "1–22 are included, while the surrounding analytical article "
                "is excluded."
            ),
        },
        {
            "id": "hin_tiwari_1995",
            "label": "Hindi · Tiwari 1995",
            "shortLabel": "Tiwari Hindi",
            "language": "Hindi",
            "languageCode": "hin",
            "script": "Devanagari",
            "role": "modern-translation",
            "color": "#9a6b32",
            "citation": (
                "Mahesh Tiwari, ed. and trans., Vijñaptimātratāsiddhiḥ: "
                "Viṃśatikā-Triṃśikābhidhāna-prakaraṇadvayātmikā, "
                "Chaukhambha Vidyabhavan, Varanasi, 1995."
            ),
            "file": (
                "Internet Archive item vijnaptimatrasiddhiofvasubandhu"
                "vigyanptimaheshtiwarichowkambhasanskritseries_492_y, "
                "printed pages 1–24"
            ),
            "rights": "rights-review",
            "rightsLabel": "Rights status requires verification",
            "extraction": (
                "Two-page spreads reconstructed from DjVu layout coordinates; "
                "Hindi translation lines were separated heuristically from the "
                "interleaved Sanskrit. OCR and segmentation require proofing."
            ),
        },
        {
            "id": "rus_lyssenko_2008",
            "label": "Russian · Lyssenko 2008",
            "shortLabel": "Lyssenko",
            "language": "Russian",
            "languageCode": "rus",
            "script": "Cyrillic",
            "role": "modern-translation",
            "color": "#7b4e68",
            "citation": (
                "В. Г. Лысенко, пер. и прим., «Вимшатика-карика-вритти. "
                "Комментарий к двадцатистишию», Вопросы философии 1 "
                "(2008), 113–131."
            ),
            "file": "50416301.pdf, PDF pages 1–16",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "Custom Cyrillic font recovered to Unicode and segmented from "
                "the numbered kārikās. Bracket glyphs and transliteration require "
                "scholarly proofing."
            ),
        },
        {
            "id": "rus_lyssenko_2022",
            "label": "Russian · Lyssenko 2022",
            "shortLabel": "Lyssenko 2022",
            "language": "Russian",
            "languageCode": "rus",
            "script": "Cyrillic",
            "role": "modern-translation",
            "color": "#86516f",
            "citation": (
                "В. Г. Лысенко, Индийские философы о природе восприятия: "
                "Дигнага и его оппоненты. Тексты и исследования, 2022."
            ),
            "file": (
                "Индийские философы о природе восприятия … (2022).djvu, "
                "OCR leaves 83–107"
            ),
            "rights": "rights-review",
            "rightsLabel": "Rights status requires verification",
            "extraction": (
                "Revised Russian translation and commentary extracted from "
                "DjVu OCR layout blocks. Translator footnotes were excluded; "
                "OCR and section boundaries require proofing."
            ),
        },
        {
            "id": "jpn_yuda_issue32",
            "label": "Japanese · Yuda",
            "shortLabel": "Yuda",
            "language": "Japanese",
            "languageCode": "jpn",
            "script": "Japanese",
            "role": "modern-translation",
            "color": "#6c5a91",
            "citation": (
                "湯田豊, 「ヴァスバンドゥの『唯識二十論』—新しい翻訳"
                "および解説」, 『法華文化研究』32."
            ),
            "file": "306_第32号_ヴァスバンドゥの「唯識二十論」 (2).pdf, pages 1–15",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "Legacy Shift-JIS text layer recovered to Unicode and segmented "
                "by semantic kārikā openings. Layout order, punctuation, and "
                "Sanskrit forms require manual Japanese scholarly proofing."
            ),
        },
        {
            "id": "jpn_yuda_watakushi_yuishiki",
            "label": "Japanese · Yuda, “Watakushi no Yuishiki”",
            "shortLabel": "Yuda essay",
            "language": "Japanese",
            "languageCode": "jpn",
            "script": "Japanese",
            "role": "modern-translation-excerpt",
            "color": "#7b5fa5",
            "citation": (
                "湯田豊, 「わたくしの“唯識” —ヴァスバンドゥの世界—」, "
                "scan supplied as 14201.pdf."
            ),
            "file": "14201.pdf, scan-only pages 1–26",
            "rights": "rights-review",
            "rightsLabel": "Rights status requires verification / OCR proofing needed",
            "extraction": (
                "Vertical Japanese scan reconstructed with Windows OCR. The "
                "source is an interpretive essay with selected Viṃśikā excerpts, "
                "not a complete 22-verse running translation."
            ),
        },
        {
            "id": "fr_levi_1932",
            "label": "French · Lévi 1932",
            "shortLabel": "Lévi",
            "language": "French",
            "languageCode": "fra",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#39765e",
            "citation": "Sylvain Lévi, Un système de philosophie bouddhique, Paris, 1932.",
            "file": "fr_levi.docx",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "DOCX OCR text; passages inferred from printed verse labels.",
        },
        {
            "id": "fra_cornu_2008",
            "label": "French · Cornu 2008",
            "shortLabel": "Cornu",
            "language": "French",
            "languageCode": "fra",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#2f7b6e",
            "citation": (
                "Philippe Cornu, trans., Cinq traités sur l’esprit seulement, "
                "Fayard, 2008."
            ),
            "file": "Cinq traités sur l’esprit seulement, EPUB chapter III",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus",
            "extraction": (
                "Structured EPUB extraction pairing each French root-verse "
                "translation with the corresponding translated auto-commentary."
            ),
        },
        {
            "id": "fra_lvp_1911",
            "label": "French · La Vallée-Poussin 1911",
            "shortLabel": "LVP French",
            "language": "French",
            "languageCode": "fra",
            "script": "Latin",
            "role": "modern-translation-ocr",
            "color": "#3d7164",
            "citation": (
                "Louis de La Vallée-Poussin, “Vasubandhu: "
                "Viṃśakakārikāprakaraṇa. Traité des vingt ślokas avec le "
                "commentaire de l’auteur,” Le Muséon 30 (1911), 53–90."
            ),
            "file": "lemuson30soc1i (1).pdf, pages 67–90",
            "rights": "public-domain",
            "rightsLabel": "Public domain / OCR proofing needed",
            "extraction": (
                "French translation segmented from numbered headings in the "
                "Le Muséon OCR layer. OCR accents and punctuation require proofing."
            ),
        },
        {
            "id": "de_frauwallner",
            "label": "German · Frauwallner",
            "shortLabel": "Frauwallner",
            "language": "German",
            "languageCode": "deu",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#526d45",
            "citation": "Erich Frauwallner, German translation, supplied PDF, pp. 366–383.",
            "file": "de_frauwallner.pdf",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "Searchable PDF text layer; passages inferred from verse labels.",
        },
        {
            "id": "de_kitayama_1934",
            "label": "German · Kitayama 1934",
            "shortLabel": "Kitayama",
            "language": "German",
            "languageCode": "deu",
            "script": "Latin",
            "role": "modern-translation-ocr",
            "color": "#6f6b38",
            "citation": (
                "Kitayama, chapter 4, “Vijñaptimātratā-siddhi-"
                "viṃśatikā,” 1934, pp. 234–268."
            ),
            "file": "User-supplied screenshots of printed pages 234–268",
            "rights": "cleared-research",
            "rightsLabel": "Cleared for research corpus / OCR proofing needed",
            "extraction": (
                "German translation reconstructed from local Windows OCR of "
                "screenshots and segmented by printed Vers headings. OCR and "
                "boundaries require proofing."
            ),
        },
    ]
    for source in sources:
        source.pop("file", None)
        source.pop("extraction", None)
    return sources


def alignment_demo() -> list[dict[str, object]]:
    return []


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source_dir", type=Path)
    parser.add_argument("output_path", type=Path)
    parser.add_argument(
        "--authorized-alignments",
        type=Path,
        help=(
            "Optional shell-native JSON exported with explicit permission from "
            "an external alignment provider."
        ),
    )
    parser.add_argument(
        "--witness-dir",
        "--english-witness-dir",
        dest="witness_dir",
        type=Path,
        default=Path(__file__).resolve().parent.parent / "source-witnesses",
        help="Directory containing the segmented witness folders.",
    )
    args = parser.parse_args()

    chinese = extract_chinese_table(args.source_dir / "zho_xuangzan.docx")
    segmented_witnesses = {
        "eng_silk_2016": load_segmented_witness_records(
            args.witness_dir / "eng_silk_2016" / "passages.json"
        ),
        "eng_siderits_2007": load_segmented_witness_records(
            args.witness_dir / "eng_siderits_2007" / "passages.json",
            require_text=False,
        ),
        "eng_kalupahana_1987": load_segmented_witness_records(
            args.witness_dir / "eng_kalupahana_1987" / "passages.json"
        ),
        "eng_anacker_2005": load_segmented_witness_records(
            args.witness_dir / "eng_anacker_2005" / "passages.json"
        ),
        "eng_tola_dragonetti_2004": load_segmented_witness_records(
            args.witness_dir / "eng_tola_dragonetti_2004" / "passages.json"
        ),
        "eng_kochumuttom_1982": load_segmented_witness_records(
            args.witness_dir / "eng_kochumuttom_1982" / "passages.json"
        ),
        "eng_wood_1991": load_segmented_witness_records(
            args.witness_dir / "eng_wood_1991" / "passages.json"
        ),
        "eng_cook_1999": load_segmented_witness_records(
            args.witness_dir / "eng_cook_1999" / "passages.json"
        ),
        "eng_hamilton_1938": load_segmented_witness_records(
            args.witness_dir / "eng_hamilton_1938" / "passages.json"
        ),
        "zho_hamilton_xuanzang_1938": load_segmented_witness_records(
            args.witness_dir
            / "zho_hamilton_xuanzang_1938"
            / "passages.json"
        ),
        "eng_cronk_1998": load_segmented_witness_records(
            args.witness_dir / "eng_cronk_1998" / "passages.json",
            require_text=False,
        ),
        "san_silk_2016": load_segmented_witness_records(
            args.witness_dir / "san_silk_2016" / "passages.json"
        ),
        "san_tola_dragonetti_2004": load_segmented_witness_records(
            args.witness_dir / "san_tola_dragonetti_2004" / "passages.json"
        ),
        "san_ruzsa_szegedi_2015": load_segmented_witness_records(
            args.witness_dir / "san_ruzsa_szegedi_2015" / "passages.json"
        ),
        "san_balcerowicz_nowakowska_1999": load_segmented_witness_records(
            args.witness_dir
            / "san_balcerowicz_nowakowska_1999"
            / "passages.json",
            require_text=False,
        ),
        "san_tiwari_1995": load_segmented_witness_records(
            args.witness_dir / "san_tiwari_1995" / "passages.json"
        ),
        "san_kalupahana_1987": load_segmented_witness_records(
            args.witness_dir / "san_kalupahana_1987" / "passages.json"
        ),
        "pol_balcerowicz_nowakowska_1999": load_segmented_witness_records(
            args.witness_dir
            / "pol_balcerowicz_nowakowska_1999"
            / "passages.json"
        ),
        "hun_szanyi_2015": load_segmented_witness_records(
            args.witness_dir / "hun_szanyi_2015" / "passages.json"
        ),
        "hin_tiwari_1995": load_segmented_witness_records(
            args.witness_dir / "hin_tiwari_1995" / "passages.json"
        ),
        "tib_silk_dunhuang_2017": load_segmented_witness_records(
            args.witness_dir / "tib_silk_dunhuang_2017" / "passages.json"
        ),
        "tib_lvp_1911": load_segmented_witness_records(
            args.witness_dir / "tib_lvp_1911" / "passages.json"
        ),
        "rus_lyssenko_2008": load_segmented_witness_records(
            args.witness_dir / "rus_lyssenko_2008" / "passages.json"
        ),
        "rus_lyssenko_2022": load_segmented_witness_records(
            args.witness_dir / "rus_lyssenko_2022" / "passages.json"
        ),
        "jpn_yuda_issue32": load_segmented_witness_records(
            args.witness_dir / "jpn_yuda_issue32" / "passages.json"
        ),
        "jpn_yuda_watakushi_yuishiki": load_segmented_witness_records(
            args.witness_dir
            / "jpn_yuda_watakushi_yuishiki"
            / "passages.json",
            require_text=False,
        ),
        "fra_cornu_2008": load_segmented_witness_records(
            args.witness_dir / "fra_cornu_2008" / "passages.json"
        ),
        "fra_lvp_1911": load_segmented_witness_records(
            args.witness_dir / "fra_lvp_1911" / "passages.json"
        ),
        "de_kitayama_1934": load_segmented_witness_records(
            args.witness_dir / "de_kitayama_1934" / "passages.json"
        ),
    }
    passages_by_source: dict[str, dict[int, str | None]] = {
        "san_levi_1925": extract_sanskrit(args.source_dir / "san_levi.docx"),
        "tib_derge": extract_tibetan(args.source_dir / "tib_derge.txt"),
        "eng_das": extract_english(args.source_dir / "eng_das.pdf"),
        "fr_levi_1932": extract_french(args.source_dir / "fr_levi.docx"),
        "de_frauwallner": extract_german(args.source_dir / "de_frauwallner.pdf"),
        **chinese,
        **{
            source_id: {
                number: records[number]["text"]
                for number in VERSE_NUMBERS
            }
            for source_id, records in segmented_witnesses.items()
        },
    }

    passages = []
    for number in VERSE_NUMBERS:
        passage_id = f"v{number}"
        passage = {
            "id": passage_id,
            "number": number,
            "label": f"Verse {number}",
            "root": SANSKRIT_ROOTS[number],
            "texts": {},
        }
        passages.append(passage)
        for source_id, source_passages in passages_by_source.items():
            text = source_passages[number] or ""
            segmented_record = segmented_witnesses.get(
                source_id,
                {},
            ).get(number, {})
            passage["texts"][source_id] = {
                "text": text,
                "tokens": tokenize_text(
                    text,
                    passage_id,
                    source_id,
                ),
                "status": segmented_record.get(
                    "status",
                    (
                        "machine-segmented"
                        if text
                        else "not-present-in-supplied-source"
                    ),
                ),
                "note": segmented_record.get(
                    "note",
                    (
                        ""
                        if text
                        else "This passage is not present in the supplied witness file."
                    ),
                ),
                "segmentationMethod": segmented_record.get("method", ""),
            }

    sources = source_records()
    for passage in passages:
        passage["sentenceUnits"] = sentence_units_for_passage(
            passage,
            sources,
        )

    alignments = [
        *alignment_demo(),
        *load_authorized_alignments(
            args.authorized_alignments,
            passages,
        ),
    ]
    passages_by_number = {passage["number"]: passage for passage in passages}
    for alignment in alignments:
        passage = passages_by_number[int(alignment["verse"])]
        alignment["targetTokenIds"] = dict(
            alignment.get("targetTokenIds", {})
        )
        for source_id, target in alignment["targets"].items():
            if alignment["targetTokenIds"].get(source_id):
                continue
            witness = passage["texts"][source_id]
            alignment["targetTokenIds"][source_id] = alignment_token_ids(
                witness["text"],
                witness["tokens"],
                target,
            )

    candidates = candidate_alignments(passages, sources, alignments)
    dharmanexus_authorized = bool(args.authorized_alignments)
    corpus = {
        "schemaVersion": "0.8.0-trial",
        "work": {
            "id": "vasubandhu-vimsika",
            "title": "Viṃśikā",
            "author": "Vasubandhu",
            "subtitle": "Complete trial parallel corpus · verses 1–22 with commentary",
        },
        "notice": (
            "Research prototype. Segmentation and phrase alignments are provisional. "
            "Corpus-wide token correspondences marked as machine-projected are low-confidence "
            "positional candidates requiring human review. "
            "Only for research purposes."
        ),
        "sources": sources,
        "passages": passages,
        "alignments": alignments,
        "candidateAlignments": candidates,
        "externalAlignmentSources": [
            {
                "id": "dharmanexus-SA_T06_vasvvmsu",
                "label": "DharmaNexus · SA_T06_vasvvmsu",
                "url": (
                    "https://dharmamitra.org/nexus/db/sa/"
                    "SA_T06_vasvvmsu/text"
                ),
                "type": "algorithmic-intertextual-matches",
                "importStatus": (
                    "authorized-reuse"
                    if dharmanexus_authorized
                    else "permission-required"
                ),
                "termsUrl": (
                    "https://dharmamitra.github.io/"
                    "dharmamitra-guides/dharmanexus/"
                ),
                "contact": "dharmamitra.project@gmail.com",
                "authorizationReference": (
                    "reference-source/dharmanexus-permission-2026-06-20.txt"
                    if dharmanexus_authorized
                    else ""
                ),
                "license": (
                    "Reuse and redistribution authorized in writing; "
                    "no separate license specified."
                    if dharmanexus_authorized
                    else ""
                ),
                "note": (
                    "Imported with written authorization received 20 June 2026. "
                    "The snapshot contains 264 Sanskrit segments and 416 Tibetan "
                    "or Chinese match records; one segment is split at a local verse "
                    "boundary, producing 265 preliminary corpus alignments."
                    if dharmanexus_authorized
                    else (
                        "The shell is ready to import an explicitly authorized export. "
                        "The public DharmaNexus API is not copied into this corpus."
                    )
                ),
            }
        ],
    }

    args.output_path.parent.mkdir(parents=True, exist_ok=True)
    args.output_path.write_text(
        json.dumps(corpus, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    javascript_corpus = copy.deepcopy(corpus)
    for passage in javascript_corpus["passages"]:
        for witness in passage["texts"].values():
            for token in witness.get("tokens", []):
                token.pop("text", None)
                token.pop("type", None)
    for candidate in javascript_corpus.get("candidateAlignments", []):
        for field in (
            "confidence",
            "level",
            "method",
            "note",
            "sourceAlignmentId",
            "status",
            "targetMethods",
        ):
            candidate.pop(field, None)
    javascript = (
        "window.CORPUS_DATA = "
        + json.dumps(javascript_corpus, ensure_ascii=False, separators=(",", ":"))
        + ";\n"
    )
    args.output_path.with_suffix(".js").write_text(javascript, encoding="utf-8")
    if args.output_path.parent.name == "data":
        (args.output_path.parent.parent / "corpus.js").write_text(
            javascript,
            encoding="utf-8",
        )


if __name__ == "__main__":
    main()
