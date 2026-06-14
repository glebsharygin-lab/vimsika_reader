from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader


VERSE_NUMBERS = list(range(1, 23))
ROMAN = {
    1: "I",
    2: "II",
    3: "III",
    4: "IV",
    5: "V",
    6: "VI",
    7: "VII",
    8: "VIII",
    9: "IX",
    10: "X",
    11: "XI",
    12: "XII",
    13: "XIII",
    14: "XIV",
    15: "XV",
    16: "XVI",
    17: "XVII",
    18: "XVIII",
    19: "XIX",
    20: "XX",
    21: "XXI",
    22: "XXII",
}
SANSKRIT_STARTS = {
    1: "vijñaptimātram evedam",
    2: "na deśakālaniyamaḥ",
    3: "deśādiniyamaḥ siddhaḥ",
    4: "svapnopaghātavat kṛtyakriyā",
    5: "tiraścāṃ saṃbhavaḥ",
    6: "yadi tatkarmabhis tatra",
    7: "karmaṇo vāsanānyatra",
    8: "rūpādyāyatanāstitvaṃ",
    9: "yataḥ svabījād vijñaptir",
    10: "tathā pudgalanairātmyapraveśo hi",
    11: "na tad ekaṃ na cānekaṃ",
    12: "ṣaṭkena yugapad yogāt",
    13: "paramāṇor asaṃyogāt",
    14: "digbhāgabhedī yasyāsti",
    15: "ekatve na krameṇetir",
    16: "pratyakṣabuddhiḥ svapnādau",
    17: "uktaṃ yathā tadābhāsā",
    18: "anyonyādhipatitvena",
    19: "maraṇam paravijñaptiviśeṣād",
    20: "kathaṃ vā daṇḍakāraṇyaśūnyatvam",
    21: "paracittavidāṃ jñānam",
    22: "vijñaptimātratāsiddhiḥ",
}
SANSKRIT_ROOTS = {
    1: (
        "vijñaptimātram evedam asadarthāvabhāsanāt |\n"
        "yadvat taimirikasyāsat keśoṇḍukādidarśanaṃ || 1 ||"
    ),
    2: (
        "na deśakālaniyamaḥ saṃtānāniyamo na ca |\n"
        "na ca kṛtyakriyā yuktā vijñaptir yadi nārthataḥ || 2 ||"
    ),
    3: (
        "deśādiniyamaḥ siddhaḥ svapnavat\n"
        "pretavat punaḥ |\n"
        "saṃtānāniyamaḥ\n"
        "sarvaiḥ pūyanadyādidarśane || 3 ||"
    ),
    4: (
        "svapnopaghātavat kṛtyakriyā\n"
        "narakavat punaḥ |\n"
        "sarvaṃ\n"
        "narakapālādidarśane taiś ca bādhane || 4 ||"
    ),
    5: (
        "tiraścāṃ saṃbhavaḥ svarge yathā na narake tathā |\n"
        "na pretānāṃ yatas tajjaṃ duḥkhaṃ nānubhavanti te || 5 ||"
    ),
    6: (
        "yadi tatkarmabhis tatra bhūtānāṃ saṃbhavas tathā |\n"
        "iṣyate pariṇāmaś ca kiṃ vijñānasya neṣyate || 6 ||"
    ),
    7: (
        "karmaṇo vāsanānyatra phalam anyatra kalpyate |\n"
        "tatraiva neṣyate yatra vāsanā kiṃ nu kāraṇaṃ || 7 ||"
    ),
    8: (
        "rūpādyāyatanāstitvaṃ tadvineyajanaṃ prati |\n"
        "abhiprāyavaśād uktam upapādukasattvavat || 8 ||"
    ),
    9: (
        "yataḥ svabījād vijñaptir yadābhāsā pravartate |\n"
        "dvividhāyatanatvena te tasyā munir abravīt || 9 ||"
    ),
    10: (
        "tathā pudgalanairātmyapraveśo hi\n"
        "anyathā punaḥ |\n"
        "deśanā dharmanairātmyapraveśaḥ\n"
        "kalpitātmanā || 10 ||"
    ),
    11: (
        "na tad ekaṃ na cānekaṃ viṣayaḥ paramāṇuśaḥ |\n"
        "na ca te saṃhatā yasmāt paramāṇur na sidhyati || 11 ||"
    ),
    12: (
        "ṣaṭkena yugapad yogāt paramāṇoḥ ṣaḍaṃśatā |\n"
        "ṣaṇṇāṃ samānadeśatvāt piṇḍaḥ syād aṇumātrakaḥ || 12 ||"
    ),
    13: (
        "paramāṇor asaṃyogāt tatsaṃghāte 'sti kasya saḥ |\n"
        "na cānavayavatvena tatsaṃyogo na sidhyati || 13 ||"
    ),
    14: (
        "digbhāgabhedī yasyāsti tasyaikatvaṃ na yujyate |\n"
        "chāyāvṛtī kathaṃ vā\n"
        "anyo na piṇḍaś cen na tasya te || 14 ||"
    ),
    15: (
        "ekatve na krameṇetir yugapan na grahāgrahau |\n"
        "vicchinnānekavṛttiś ca sūkṣmānīkṣā ca no bhavet || 15 ||"
    ),
    16: (
        "pratyakṣabuddhiḥ svapnādau yathā\n"
        "sā ca yadā tadā |\n"
        "na so 'rtho dṛśyate tasya pratyakṣatvaṃ kathaṃ mataṃ || 16 ||"
    ),
    17: (
        "uktaṃ yathā tadābhāsā vijñaptiḥ\n"
        "smaraṇaṃ tataḥ |\n"
        "svapnadṛgviṣayābhāvaṃ nāprabuddho 'vigacchati || 17 ||"
    ),
    18: (
        "anyonyādhipatitvena vijñaptiniyamo mithaḥ |\n"
        "middhenopahataṃ cittaṃ svapne tenāsamaṃ phalaṃ || 18 ||"
    ),
    19: (
        "maraṇam paravijñaptiviśeṣād vikriyā yathā |\n"
        "smṛtilopādikānyeṣāṃ piśācādimanovaśāt || 19 ||"
    ),
    20: (
        "kathaṃ vā daṇḍakāraṇyaśūnyatvam ṛṣikopataḥ |\n"
        "manodaṇḍo mahāvadyaḥ kathaṃ vā tena sidhyati || 20 ||"
    ),
    21: (
        "paracittavidāṃ jñānam ayathārthaṃ kathaṃ yathā |\n"
        "svacittajñānaṃ\n"
        "ajñānād yathā buddhasya gocaraḥ || 21 ||"
    ),
    22: (
        "vijñaptimātratāsiddhiḥ svaśaktisadṛśī mayā |\n"
        "kṛteyaṃ sarvathā sā tu na cintyā\n"
        "buddhagocaraḥ || 22 ||"
    ),
}


def read_docx_paragraphs(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs if paragraph.text.strip()]


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
    text = re.sub(r"(?<=\w)-\s*\n\s*(?=\w)", "", text)
    return re.sub(r"[ \t]+", " ", text)


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def slice_by_markers(
    text: str,
    markers: dict[int, re.Pattern[str]],
    first_start: int | None = None,
) -> dict[int, str]:
    positions: dict[int, int] = {}
    cursor = 0
    for number in VERSE_NUMBERS:
        match = markers[number].search(text, cursor)
        if not match:
            raise ValueError(f"Could not locate verse {number}")
        positions[number] = match.start()
        cursor = match.start() + 1

    if first_start is not None:
        positions[VERSE_NUMBERS[0]] = first_start

    result: dict[int, str] = {}
    for number in VERSE_NUMBERS:
        start = positions[number]
        end = positions.get(number + 1, len(text))
        result[number] = clean_text(text[start:end])
    return result


def extract_sanskrit(path: Path) -> dict[int, str]:
    text = "\n\n".join(read_docx_paragraphs(path))
    markers = {
        number: re.compile(re.escape(start), re.IGNORECASE)
        for number, start in SANSKRIT_STARTS.items()
    }
    return slice_by_markers(text, markers, first_start=0)


def extract_tibetan(path: Path) -> dict[int, str]:
    text = path.read_text(encoding="utf-8")
    markers = {
        number: re.compile(rf"(?m)^{ROMAN[number]}\s*$")
        for number in VERSE_NUMBERS
    }
    return slice_by_markers(text, markers, first_start=0)


def extract_english(path: Path) -> dict[int, str | None]:
    text = read_pdf(path)
    available_numbers = [1, 2, 3, 4, *range(11, 23)]
    markers = {
        number: re.compile(rf"\bVerse\s+{number}\.", re.IGNORECASE)
        for number in available_numbers
    }
    positions: dict[int, int] = {}
    cursor = 0
    for number in available_numbers:
        match = markers[number].search(text, cursor)
        if not match:
            raise ValueError(f"Could not locate English verse {number}")
        positions[number] = match.start()
        cursor = match.start() + 1

    statement_start = text.find("I. A Statement of the View")
    if statement_start >= 0:
        positions[1] = statement_start

    passages: dict[int, str | None] = {number: None for number in VERSE_NUMBERS}
    for index, number in enumerate(available_numbers):
        start = positions[number]
        end = (
            positions[available_numbers[index + 1]]
            if index + 1 < len(available_numbers)
            else len(text)
        )
        passages[number] = clean_text(text[start:end])
    return passages


def extract_german(path: Path) -> dict[int, str]:
    text = read_pdf(path)
    markers = {
        number: re.compile(rf"\bV\.\s*{number}\b", re.IGNORECASE)
        for number in VERSE_NUMBERS
    }
    opening = text.find("KAPITEL")
    return slice_by_markers(text, markers, first_start=max(opening, 0))


def extract_french(path: Path) -> dict[int, str]:
    text = clean_text("\n\n".join(read_docx_paragraphs(path)))
    starts = {
        1: "Le Grand Véhicule établit",
        2: "2. – Si la Notification",
        3: "3 a. La détermination",
        4: "4 a. Pour l'accomplissement",
        5: "5. S'il peut",
        6: "6. Si vous admettez",
        7: "7. Vous imaginez",
        8: "8. S'il est parlé",
        9: "Le germe propre d'où la Notification",
        10: "C'est que l'entrée dans le Sans-Soi",
        11: "11. Le Domaine",
        12: "12 ab. Si l'Atome",
        13: "13 ab. – Si l'Atome",
        14: "14 ab. Du moment",
        15: "15. Dans",
        16: "16 a. L'idée",
        17: "17 a. Nous avons montré",
        18: "18 ab. C'est par",
        19: "19. La mort",
        20: "20 ab. – Ou bien",
        21: "21 abc. Chez ceux",
        22: "22 abc. Cette démonstration",
    }
    positions: dict[int, int] = {}
    cursor = 0
    for number in VERSE_NUMBERS:
        position = text.find(starts[number], cursor)
        if position < 0:
            raise ValueError(f"Could not locate French verse {number}: {starts[number]}")
        positions[number] = position
        cursor = position + 1

    return {
        number: clean_text(
            text[positions[number] : positions.get(number + 1, len(text))]
        )
        for number in VERSE_NUMBERS
    }


def cell_has_verse(cell: str, number: int) -> bool:
    return bool(re.search(rf"\(\s*{number}(?:[a-d-]*)?\s*\)", cell, re.IGNORECASE))


def extract_chinese_table(path: Path) -> dict[str, dict[int, str]]:
    document = Document(path)
    table = document.tables[0]
    rows = [[clean_text(cell.text) for cell in row.cells] for row in table.rows]
    columns = {
        "zho_xuanzang": 0,
        "zho_paramartha": 1,
        "zho_bodhiruci": 2,
    }

    starts: dict[int, int] = {}
    cursor = 0
    for number in VERSE_NUMBERS:
        for row_index in range(cursor, len(rows)):
            if cell_has_verse(rows[row_index][3], number):
                starts[number] = row_index
                cursor = row_index + 1
                break
        else:
            raise ValueError(f"Could not locate Chinese table verse {number}")
    starts[1] = 1

    passages = {version_id: {} for version_id in columns}
    for number in VERSE_NUMBERS:
        start = starts[number]
        end = starts.get(number + 1, len(rows))
        passage_rows = rows[start:end]
        for version_id, column_index in columns.items():
            values = [row[column_index] for row in passage_rows if row[column_index]]
            passages[version_id][number] = clean_text("\n\n".join(values))

    return passages


def source_records() -> list[dict[str, object]]:
    return [
        {
            "id": "san_levi_1925",
            "label": "Sanskrit · Lévi 1925",
            "shortLabel": "Sanskrit",
            "language": "Sanskrit",
            "languageCode": "san",
            "script": "Latin transliteration",
            "role": "edition",
            "color": "#a64b2a",
            "citation": "Sylvain Lévi (ed.), Vijñaptimātratāsiddhi, Paris, 1925.",
            "file": "san_levi.docx",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "DOCX text layer; passage boundaries inferred from verse openings.",
        },
        {
            "id": "tib_derge",
            "label": "Tibetan · Derge",
            "shortLabel": "Tibetan",
            "language": "Tibetan",
            "languageCode": "bod",
            "script": "Wylie transliteration",
            "role": "canonical-translation",
            "color": "#8b6a20",
            "citation": "Nyi shu pa’i ’grel pa, Derge Bstan ’gyur tradition.",
            "file": "tib_derge.txt",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "UTF-8 text; sections inferred from Roman-numeral headings.",
        },
        {
            "id": "zho_xuanzang",
            "label": "Chinese · Xuanzang",
            "shortLabel": "Xuanzang",
            "language": "Chinese",
            "languageCode": "zho",
            "script": "Traditional Chinese",
            "role": "canonical-translation",
            "color": "#9a3f57",
            "citation": "唯識二十論, translated by Xuanzang, Taishō T1590.",
            "file": "zho_xuangzan.docx",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "First column of the supplied four-column comparative table.",
        },
        {
            "id": "zho_paramartha",
            "label": "Chinese · Paramārtha",
            "shortLabel": "Paramārtha",
            "language": "Chinese",
            "languageCode": "zho",
            "script": "Traditional Chinese",
            "role": "canonical-translation",
            "color": "#7f4a77",
            "citation": "大乘唯識論, translated by Paramārtha, Taishō T1589.",
            "file": "zho_xuangzan.docx",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "Second column of the supplied four-column comparative table.",
        },
        {
            "id": "zho_bodhiruci",
            "label": "Chinese · Bodhiruci",
            "shortLabel": "Bodhiruci",
            "language": "Chinese",
            "languageCode": "zho",
            "script": "Traditional Chinese",
            "role": "canonical-translation",
            "color": "#6d5a8d",
            "citation": "唯識論, translated by Bodhiruci, Taishō T1588.",
            "file": "zho_xuangzan.docx",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "Third column of the supplied four-column comparative table.",
        },
        {
            "id": "eng_das",
            "label": "English · Nilanjan Das",
            "shortLabel": "Das",
            "language": "English",
            "languageCode": "eng",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#2d6f8f",
            "citation": "Nilanjan Das, Twenty Verses with Auto-Commentary, online draft.",
            "file": "eng_das.pdf",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "Searchable PDF text layer; passages inferred from verse labels.",
        },
        {
            "id": "fr_levi_1932",
            "label": "French · Lévi 1932",
            "shortLabel": "Lévi",
            "language": "French",
            "languageCode": "fra",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#39765e",
            "citation": "Sylvain Lévi, Un système de philosophie bouddhique, Paris, 1932.",
            "file": "fr_levi.docx",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "DOCX OCR text; passages inferred from printed verse labels.",
        },
        {
            "id": "de_frauwallner",
            "label": "German · Frauwallner",
            "shortLabel": "Frauwallner",
            "language": "German",
            "languageCode": "deu",
            "script": "Latin",
            "role": "modern-translation",
            "color": "#526d45",
            "citation": "Erich Frauwallner, German translation, supplied PDF, pp. 366–383.",
            "file": "de_frauwallner.pdf",
            "rights": "public-domain",
            "rightsLabel": "Public domain / open access",
            "extraction": "Searchable PDF text layer; passages inferred from verse labels.",
        },
    ]


def alignment_demo() -> list[dict[str, object]]:
    return [
        {
            "id": "v15-unity",
            "verse": 15,
            "label": "unity / one entity",
            "note": "Provisional phrase-level equivalence; requires philological review.",
            "targets": {
                "san_levi_1925": "ekatve",
                "tib_derge": "gcig na",
                "zho_xuanzang": "一應無",
                "zho_paramartha": "若一",
                "zho_bodhiruci": "若一",
                "eng_das": "If it were simple",
                "fr_levi_1932": "l'unité",
                "de_frauwallner": "Im Falle der Einheit",
            },
        },
        {
            "id": "v15-traversal",
            "verse": 15,
            "label": "gradual traversal",
            "note": "A compact concept alignment, not a claim of literal word identity.",
            "targets": {
                "san_levi_1925": "krameṇetir",
                "tib_derge": "rim gyis ’gro ba",
                "zho_xuanzang": "次行",
                "zho_paramartha": "次行",
                "zho_bodhiruci": "行不次",
                "eng_das": "gradual traversal",
                "fr_levi_1932": "marcher au pas",
                "de_frauwallner": "schrittweise Bewegung",
            },
        },
        {
            "id": "v15-apprehension",
            "verse": 15,
            "label": "apprehension / non-apprehension",
            "note": "The translations distribute the contrast across different syntactic forms.",
            "targets": {
                "san_levi_1925": "grahāgrahau",
                "tib_derge": "zin dang ma zin",
                "zho_xuanzang": "至未至",
                "zho_paramartha": "已未得",
                "zho_bodhiruci": "取捨",
                "eng_das": "perception and non-perception",
                "fr_levi_1932": "prendre et ne pas prendre",
                "de_frauwallner": "Erfassen und Nichterfassen",
            },
        },
        {
            "id": "v15-subtle",
            "verse": 15,
            "label": "the very small / subtle",
            "note": "This alignment exposes a useful difference between nominal and verbal renderings.",
            "targets": {
                "san_levi_1925": "sūkṣmānīkṣā",
                "tib_derge": "phra ba",
                "zho_xuanzang": "難見細物",
                "zho_paramartha": "細難見",
                "zho_bodhiruci": "微細亦應見",
                "eng_das": "very small objects be imperceptible",
                "fr_levi_1932": "voir l'infiniment petit",
                "de_frauwallner": "Nichtwahrnehmung",
            },
        },
    ]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source_dir", type=Path)
    parser.add_argument("output_path", type=Path)
    args = parser.parse_args()

    chinese = extract_chinese_table(args.source_dir / "zho_xuangzan.docx")
    passages_by_source: dict[str, dict[int, str | None]] = {
        "san_levi_1925": extract_sanskrit(args.source_dir / "san_levi.docx"),
        "tib_derge": extract_tibetan(args.source_dir / "tib_derge.txt"),
        "eng_das": extract_english(args.source_dir / "eng_das.pdf"),
        "fr_levi_1932": extract_french(args.source_dir / "fr_levi.docx"),
        "de_frauwallner": extract_german(args.source_dir / "de_frauwallner.pdf"),
        **chinese,
    }

    passages = []
    for number in VERSE_NUMBERS:
        passages.append(
            {
                "id": f"v{number}",
                "number": number,
                "label": f"Verse {number}",
                "root": SANSKRIT_ROOTS[number],
                "texts": {
                    source_id: {
                        "text": source_passages[number] or "",
                        "status": (
                            "machine-segmented"
                            if source_passages[number]
                            else "not-present-in-supplied-source"
                        ),
                        "note": (
                            ""
                            if source_passages[number]
                            else "This passage is not present in the supplied witness file."
                        ),
                    }
                    for source_id, source_passages in passages_by_source.items()
                },
            }
        )

    corpus = {
        "schemaVersion": "0.2.0-trial",
        "work": {
            "id": "vasubandhu-vimsika",
            "title": "Viṃśikā",
            "author": "Vasubandhu",
            "subtitle": "Complete trial parallel corpus · verses 1–22 with commentary",
        },
        "notice": (
            "Scholarly prototype. Segmentation and phrase alignments are provisional. "
            "The project owner has confirmed the supplied witnesses are cleared for public scholarly publication."
        ),
        "sources": source_records(),
        "passages": passages,
        "alignments": alignment_demo(),
    }

    args.output_path.parent.mkdir(parents=True, exist_ok=True)
    args.output_path.write_text(
        json.dumps(corpus, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    javascript_path = args.output_path.with_suffix(".js")
    javascript_path.write_text(
        "window.CORPUS_DATA = "
        + json.dumps(corpus, ensure_ascii=False, indent=2)
        + ";\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
