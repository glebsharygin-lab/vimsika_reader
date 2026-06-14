from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def extract_docx(path: Path) -> dict[str, object]:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]
    return {
        "kind": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
    }


def extract_pdf(path: Path) -> dict[str, object]:
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return {
        "kind": "pdf",
        "pages": pages,
    }


def extract_txt(path: Path) -> dict[str, object]:
    return {
        "kind": "txt",
        "text": path.read_text(encoding="utf-8"),
    }


def extract(path: Path) -> dict[str, object]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".txt":
        return extract_txt(path)
    raise ValueError(f"Unsupported source format: {path}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source_dir", type=Path)
    parser.add_argument("output_dir", type=Path)
    args = parser.parse_args()

    args.output_dir.mkdir(parents=True, exist_ok=True)
    manifest = []

    for path in sorted(args.source_dir.iterdir()):
        if not path.is_file():
            continue
        try:
            result = extract(path)
        except ValueError:
            continue

        output_path = args.output_dir / f"{path.stem}.json"
        output_path.write_text(
            json.dumps(result, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        manifest.append(
            {
                "source": str(path),
                "output": str(output_path),
                "kind": result["kind"],
                "size": path.stat().st_size,
            }
        )

    (args.output_dir / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
